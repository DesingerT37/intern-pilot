"""
导出服务模块
负责将 Markdown 格式的简历转换为 PDF 或 DOCX 格式
"""
import os
import re
from pathlib import Path
import markdown
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from loguru import logger


# 常见中文字体路径（Windows / Linux）
_CHINESE_FONT_CANDIDATES = [
    Path(r"C:\Windows\Fonts\msyh.ttc"),       # 微软雅黑
    Path(r"C:\Windows\Fonts\msyhbd.ttc"),
    Path(r"C:\Windows\Fonts\simhei.ttf"),     # 黑体
    Path(r"C:\Windows\Fonts\simsun.ttc"),       # 宋体
    Path("/usr/share/fonts/truetype/wqy/wqy-microhei.ttc"),
    Path("/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc"),
    Path("/System/Library/Fonts/PingFang.ttc"),
]


class ExportService:
    """导出服务类"""

    # 样式模板定义（PDF 使用 CSS 值，DOCX 使用 docx_* 字段）
    STYLES = {
        "default": {
            "font_family": "'Arial', 'Microsoft YaHei', 'SimHei', sans-serif",
            "font_size": "11pt",
            "line_height": "1.6",
            "color": "#333",
            "h1_size": "24pt",
            "h1_color": "#000000",
            "h1_border": "2px solid #333",
            "h2_size": "16pt",
            "h2_color": "#2c3e50",
            "h3_size": "13pt",
            "h3_color": "#34495e",
            "docx_body_font": "Microsoft YaHei",
            "docx_body_size": 11,
            "docx_h1_rgb": (0, 0, 0),
            "docx_h2_rgb": (44, 62, 80),
            "docx_h3_rgb": (52, 73, 94),
        },
        "modern": {
            "font_family": "'Helvetica', 'Microsoft YaHei', 'SimHei', sans-serif",
            "font_size": "10.5pt",
            "line_height": "1.5",
            "color": "#2c3e50",
            "h1_size": "22pt",
            "h1_color": "#3498db",
            "h1_border": "3px solid #3498db",
            "h2_size": "15pt",
            "h2_color": "#2980b9",
            "h3_size": "12pt",
            "h3_color": "#34495e",
            "docx_body_font": "Microsoft YaHei",
            "docx_body_size": 10,
            "docx_h1_rgb": (52, 152, 219),
            "docx_h2_rgb": (41, 128, 185),
            "docx_h3_rgb": (52, 73, 94),
        },
        "classic": {
            "font_family": "'Times New Roman', 'SimSun', serif",
            "font_size": "12pt",
            "line_height": "1.8",
            "color": "#000000",
            "h1_size": "26pt",
            "h1_color": "#000000",
            "h1_border": "1px solid #000000",
            "h2_size": "18pt",
            "h2_color": "#000000",
            "h3_size": "14pt",
            "h3_color": "#333333",
            "docx_body_font": "SimSun",
            "docx_body_size": 12,
            "docx_h1_rgb": (0, 0, 0),
            "docx_h2_rgb": (0, 0, 0),
            "docx_h3_rgb": (51, 51, 51),
        },
    }

    @staticmethod
    def _resolve_style(style: str) -> dict:
        return ExportService.STYLES.get(style, ExportService.STYLES["default"])

    @staticmethod
    def _ensure_parent_dir(output_path: str) -> None:
        parent = os.path.dirname(os.path.abspath(output_path))
        if parent:
            os.makedirs(parent, exist_ok=True)

    @staticmethod
    def _build_font_face_css() -> str:
        """为 PDF 生成 @font-face，确保 WeasyPrint 能加载中文字体"""
        blocks = []
        for font_path in _CHINESE_FONT_CANDIDATES:
            if not font_path.is_file():
                continue
            uri = font_path.as_uri()
            family = "ResumeChinese"
            blocks.append(
                f"""
        @font-face {{
            font-family: '{family}';
            src: url('{uri}');
        }}"""
            )
            break
        return "".join(blocks)

    @staticmethod
    def markdown_to_html(markdown_content: str) -> str:
        """
        将 Markdown 转换为 HTML

        Args:
            markdown_content: Markdown 内容

        Returns:
            str: HTML 内容（body 片段，不含完整文档）
        """
        try:
            html_content = markdown.markdown(
                markdown_content,
                extensions=[
                    "extra",
                    "codehilite",
                    "nl2br",
                    "sane_lists",
                ],
            )
            logger.info("Markdown 转 HTML 成功")
            return html_content
        except Exception as e:
            logger.error(f"Markdown 转 HTML 失败: {e}")
            raise ValueError(f"Markdown 转换失败: {e}") from e

    @staticmethod
    def apply_pdf_styles(html_content: str, style: str = "default") -> str:
        """
        应用 PDF 样式

        Args:
            html_content: HTML 内容
            style: 样式模板名称（default/modern/classic）

        Returns:
            str: 带样式的完整 HTML 文档
        """
        style_config = ExportService._resolve_style(style)
        font_face_css = ExportService._build_font_face_css()
        chinese_fallback = (
            "'ResumeChinese', 'Microsoft YaHei', 'SimHei', 'SimSun', "
            "'Noto Sans CJK SC', 'WenQuanYi Micro Hei', sans-serif"
        )
        font_family = f"{chinese_fallback}, {style_config['font_family']}"

        styled_html = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <style>
        {font_face_css}

        @page {{
            size: A4;
            margin: 2cm;
        }}

        body {{
            font-family: {font_family};
            font-size: {style_config['font_size']};
            line-height: {style_config['line_height']};
            color: {style_config['color']};
            max-width: 100%;
        }}

        h1 {{
            font-size: {style_config['h1_size']};
            color: {style_config['h1_color']};
            margin-bottom: 10px;
            padding-bottom: 8px;
            border-bottom: {style_config['h1_border']};
            page-break-after: avoid;
        }}

        h2 {{
            font-size: {style_config['h2_size']};
            color: {style_config['h2_color']};
            margin-top: 20px;
            margin-bottom: 10px;
            page-break-after: avoid;
        }}

        h3 {{
            font-size: {style_config['h3_size']};
            color: {style_config['h3_color']};
            margin-top: 15px;
            margin-bottom: 8px;
            page-break-after: avoid;
        }}

        p {{
            margin-bottom: 8px;
            text-align: justify;
        }}

        ul, ol {{
            margin-left: 20px;
            margin-bottom: 10px;
        }}

        li {{
            margin-bottom: 5px;
        }}

        strong {{
            font-weight: bold;
            color: {style_config['h2_color']};
        }}

        em {{
            font-style: italic;
        }}

        code {{
            background-color: #f4f4f4;
            padding: 2px 4px;
            border-radius: 3px;
            font-family: 'Courier New', monospace;
            font-size: 0.9em;
        }}

        pre {{
            background-color: #f4f4f4;
            padding: 10px;
            border-radius: 5px;
            overflow-x: auto;
            page-break-inside: avoid;
        }}

        pre code {{
            background-color: transparent;
            padding: 0;
        }}

        table {{
            border-collapse: collapse;
            width: 100%;
            margin-bottom: 15px;
        }}

        th, td {{
            border: 1px solid #ddd;
            padding: 8px;
            text-align: left;
        }}

        th {{
            background-color: #f2f2f2;
            font-weight: bold;
        }}

        a {{
            color: #3498db;
            text-decoration: none;
        }}

        blockquote {{
            border-left: 4px solid #ddd;
            padding-left: 15px;
            margin-left: 0;
            color: #666;
            font-style: italic;
        }}

        hr {{
            border: none;
            border-top: 1px solid #ddd;
            margin: 20px 0;
        }}
    </style>
</head>
<body>
    {html_content}
</body>
</html>"""

        logger.info(f"应用 PDF 样式: {style}")
        return styled_html

    @staticmethod
    def export_to_pdf(
        markdown_content: str,
        output_path: str,
        style: str = "default",
    ) -> str:
        """
        导出为 PDF

        Args:
            markdown_content: Markdown 内容
            output_path: 输出文件路径
            style: 样式模板（default/modern/classic）

        Returns:
            str: 生成的文件路径
        """
        try:
            from weasyprint import HTML
            from weasyprint.text.fonts import FontConfiguration
        except (ImportError, OSError) as e:
            logger.error(f"WeasyPrint 不可用: {e}")
            raise RuntimeError(
                "PDF 导出需要 WeasyPrint 及其系统依赖（GTK/Pango）。"
                "请参考 https://doc.courtbouillon.org/weasyprint/stable/first_steps.html"
            ) from e

        try:
            logger.info(f"开始导出 PDF: {output_path}")
            ExportService._ensure_parent_dir(output_path)

            html_content = ExportService.markdown_to_html(markdown_content)
            styled_html = ExportService.apply_pdf_styles(html_content, style)

            font_config = FontConfiguration()
            HTML(string=styled_html).write_pdf(output_path, font_config=font_config)

            logger.info(f"PDF 导出成功: {output_path}")
            return output_path
        except RuntimeError:
            raise
        except Exception as e:
            logger.error(f"PDF 导出失败: {e}")
            raise ValueError(f"PDF 导出失败: {e}") from e

    @staticmethod
    def export_to_docx(
        markdown_content: str,
        output_path: str,
        style: str = "default",
    ) -> str:
        """
        导出为 DOCX

        Args:
            markdown_content: Markdown 内容
            output_path: 输出文件路径
            style: 样式模板（default/modern/classic）

        Returns:
            str: 生成的文件路径
        """
        try:
            logger.info(f"开始导出 DOCX: {output_path}")
            ExportService._ensure_parent_dir(output_path)

            doc = Document()
            ExportService._apply_docx_styles(doc, style)

            lines = markdown_content.split("\n")
            i = 0
            while i < len(lines):
                line = lines[i]

                if line.startswith("# "):
                    heading = doc.add_heading(line[2:].strip(), level=1)
                    ExportService._style_heading(heading, 1, style)
                elif line.startswith("## "):
                    heading = doc.add_heading(line[3:].strip(), level=2)
                    ExportService._style_heading(heading, 2, style)
                elif line.startswith("### "):
                    heading = doc.add_heading(line[4:].strip(), level=3)
                    ExportService._style_heading(heading, 3, style)
                elif line.startswith("- ") or line.startswith("* "):
                    text = line[2:].strip()
                    p = doc.add_paragraph(style="List Bullet")
                    ExportService._add_formatted_text(p, text, style)
                elif re.match(r"^\d+\.\s", line):
                    text = re.sub(r"^\d+\.\s", "", line).strip()
                    p = doc.add_paragraph(style="List Number")
                    ExportService._add_formatted_text(p, text, style)
                elif line.startswith("```"):
                    code_lines = []
                    i += 1
                    while i < len(lines) and not lines[i].startswith("```"):
                        code_lines.append(lines[i])
                        i += 1
                    code_text = "\n".join(code_lines)
                    p = doc.add_paragraph()
                    run = p.add_run(code_text)
                    run.font.name = "Courier New"
                    run.font.size = Pt(9)
                    ExportService._set_run_east_asia_font(run, "Courier New")
                elif line.strip():
                    p = doc.add_paragraph()
                    ExportService._add_formatted_text(p, line.strip(), style)

                i += 1

            doc.save(output_path)
            logger.info(f"DOCX 导出成功: {output_path}")
            return output_path
        except Exception as e:
            logger.error(f"DOCX 导出失败: {e}")
            raise ValueError(f"DOCX 导出失败: {e}") from e

    @staticmethod
    def _apply_docx_styles(doc: Document, style: str) -> None:
        """应用 DOCX 文档默认样式（含中文字体）"""
        style_config = ExportService._resolve_style(style)
        normal_style = doc.styles["Normal"]
        normal_font = normal_style.font
        normal_font.name = style_config["docx_body_font"]
        normal_font.size = Pt(style_config["docx_body_size"])
        ExportService._set_style_east_asia_font(normal_style, style_config["docx_body_font"])

    @staticmethod
    def _set_style_east_asia_font(style, font_name: str) -> None:
        rpr = style.element.rPr
        if rpr is not None and rpr.rFonts is not None:
            rpr.rFonts.set(qn("w:eastAsia"), font_name)

    @staticmethod
    def _set_run_east_asia_font(run, font_name: str) -> None:
        run.font.name = font_name
        rpr = run._element.get_or_add_rPr()
        rfonts = rpr.get_or_add_rFonts()
        rfonts.set(qn("w:eastAsia"), font_name)

    @staticmethod
    def _style_heading(heading, level: int, style: str) -> None:
        """按模板设置标题样式"""
        style_config = ExportService._resolve_style(style)
        size_map = {1: 24, 2: 16, 3: 13}
        color_map = {
            1: style_config["docx_h1_rgb"],
            2: style_config["docx_h2_rgb"],
            3: style_config["docx_h3_rgb"],
        }
        size = size_map.get(level, 13)
        rgb = color_map.get(level, (0, 0, 0))
        font_name = style_config["docx_body_font"]

        for run in heading.runs:
            run.font.size = Pt(size)
            run.font.color.rgb = RGBColor(*rgb)
            run.font.bold = True
            ExportService._set_run_east_asia_font(run, font_name)

    @staticmethod
    def _add_formatted_text(paragraph, text: str, style: str = "default") -> None:
        """
        添加格式化文本（支持加粗、斜体、行内代码）

        Args:
            paragraph: docx 段落对象
            text: 文本内容
            style: 样式模板名称
        """
        style_config = ExportService._resolve_style(style)
        font_name = style_config["docx_body_font"]
        # 按 token 切分：**bold** *italic* `code`
        pattern = r"(\*\*.+?\*\*|__.+?__|\*.+?\*|_.+?_|`.+?`)"
        parts = re.split(pattern, text)

        for part in parts:
            if not part:
                continue
            if part.startswith("**") and part.endswith("**"):
                run = paragraph.add_run(part[2:-2])
                run.font.bold = True
            elif part.startswith("__") and part.endswith("__"):
                run = paragraph.add_run(part[2:-2])
                run.font.bold = True
            elif (
                (part.startswith("*") and part.endswith("*") and not part.startswith("**"))
                or (part.startswith("_") and part.endswith("_") and not part.startswith("__"))
            ):
                run = paragraph.add_run(part[1:-1])
                run.font.italic = True
            elif part.startswith("`") and part.endswith("`"):
                run = paragraph.add_run(part[1:-1])
                run.font.name = "Courier New"
                run.font.size = Pt(9)
                ExportService._set_run_east_asia_font(run, "Courier New")
                continue
            else:
                run = paragraph.add_run(part)

            ExportService._set_run_east_asia_font(run, font_name)


export_service = ExportService()
